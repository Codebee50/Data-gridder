from django.shortcuts import render, redirect
from .models import Profile, Form, FormValue,Contact
from django.contrib.auth.models import User
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, FileResponse
import os
import json
from django.core.serializers import serialize
from docx import Document
import shutil
from django.core.mail import send_mail
from django.conf import settings
from . import forms as mForms
from django.template.loader import render_to_string
from django.contrib.sites.shortcuts import get_current_site
from functools import partial
from django.core.exceptions import ObjectDoesNotExist
from django.core.files import File

# Create your views here.
# @login_required(login_url='register')
def home(request):
    if request.user.is_authenticated:
        #this means that the user is logged in 
        user_object = User.objects.get(id= request.user.id)
        user_profile = Profile.objects.get(user=user_object)

        # current_user = request.user.username
        contact_form = mForms.ContactForm()
        context= {
            'user_exists': 'true',
            'user_profile' : user_profile,
            'contact_form': contact_form,
        }
        return render(request, 'index.html', context)
    else:
        #this means that the user is not logged in 
        contact_form = mForms.ContactForm()
        context = {
            'user_exists': 'false',
            'contact_form': contact_form,
        }
        return render(request, 'index.html', context)

#This is used to pass in the nessesary requirements for displaying the dashboard screen        
@login_required(login_url='login')
def dashboard(request):
    try:
        user_object = User.objects.get(id= request.user.id)
        user_profile = Profile.objects.get(user=user_object)
    except ObjectDoesNotExist:
        user_object = None
        user_profile = None
    
    current_user = request.user.username

    forms = Form.objects.filter(form_author = current_user)
    registered_forms = FormValue.objects.filter(user_name = current_user).order_by('-registered_date')
    contact_form = mForms.ContactForm()
    try:
        form_values_list = list(forms.values())
        registered_form_list = list(registered_forms.values())
    except Exception as e:
        print(e)
        form_values_list = []
        registered_form_list = []
    context= {
        'user_profile' : user_profile,
        'forms': form_values_list,
        'registered_forms': registered_form_list,
        'contact_form': contact_form
    }
    return render(request, 'dashboard.html', context)

def getuserforms(reqeust):
    # for form in Form.objects.filter(form_creator__isnull = True):
    #     try:
    #         user = User.objects.get(username = form.form_author)
    #         form.form_creator = user
    #         form.save()
    #         print(f'Creator updated for {form.form_author}')
    #     except User.DoesNotExist:
    #         pass

    # print('All creators updated successfully')

    user_id = reqeust.user.id#getting the id of the user 
    try: 
        user_object = User.objects.get(id = user_id)
    except User.DoesNotExist:
        return JsonResponse({
            'status': 500,
            'message': 'User does not exist'
        }, status=500)
    
    forms = Form.objects.filter(form_creator = user_object)

    try: 
        form_list = list(forms.values())
    except Exception as e:
        print(e)
        form_list = None

    context = {
        'status': 200, 
        'forms': form_list
    }
    return JsonResponse(context, status=200)



"""this function checks if the form exists and if the current user is authenticated 
    and returns json responses with messages matching the corresponding cases"""
def findform(request):
    if request.method == 'POST':
        formcode = request.POST['formcode']
        if Form.objects.filter(form_code=formcode).exists():
            if request.user.is_authenticated:      
                form = Form.objects.get(form_code=formcode)
                context = {
                    'status': 'success',
                    'message': 'Form exists',
                    'formname': form.form_name,
                    'formauthor': form.form_author,
                    'formcode': form.form_code

                }
                return JsonResponse(context)
            else:
                context = {
                    'status': 'error',
                    'message': 'un_auth'
                }
                return JsonResponse(context)
        else:
            context = {
                'status': 'error',
                'message': 'n_f'
            }
            return JsonResponse(context)

    else:
        return render(request, '/')


""" gets the fields of the form with the provided form code
     also gets the form value withe the provided variable pk if user wants to edit his entry
     returns a json response"""
def getform(request, formcode, pk):
    if request.method == 'GET':
        if pk == 'nb':
            form = Form.objects.get(form_code=formcode)
            json_fields = json.dumps(form.fields)
            context = {
                'fields': json_fields,
                'values': 'empty'
            }
            
            return JsonResponse(context)

        else:
            if FormValue.objects.filter(id = pk).exists():
                values = FormValue.objects.get(id = pk)
                form = Form.objects.get(form_code=formcode)
                json_fields = json.dumps(form.fields)
                serialized_values = serialize('json', [values])
                context = {
                    'fields': json_fields,
                    'values': serialized_values
                }
                return JsonResponse(context)
               
            else:
                context = {
                    'fields': 'empty',
                    'values':  'empty',
                    'message': 'Form does not exist'
                }
                return JsonResponse(context)
    else: 
        pass
    
# @login_required(login_url='login')
def registerForm(request, formcode, pk):
    #checking if form exists
    guest = request.GET.get('guest')
    if Form.objects.filter(form_code= formcode).exists():
        form = Form.objects.get(form_code = formcode)
        if form.status == Form.Status.LOCKED:
            return render(request, 'form-locked.html')
        formname = form.form_name
        formauthor = form.form_author
        formvalues = FormValue.objects.filter(form_code = formcode)
        itemcount = formvalues.count()
        
        context = {
            'stautus' : 'success',
            'statusCode': 200,
            'formcode': formcode,
            'formname': formname,
            'formauthor': formauthor,
            'itemcount': itemcount,
            'value_id': pk,
            'is_authenticated': request.user.is_authenticated,
            'guest': guest, 
            'description': 'Please fill the form below' if form.description == None else form.description
        }
        return render(request, 'regform.html', context)
    else:
        return render(request, 'form-not-found.html')
        # context = {
        #     'status': 'failed',
        #     'statusCode': 401,
        #     'message': 'Form ' + formcode + ' does not exist',
        #     'formcode': 'none',
        #     'formauthor': 'none',
        #     'formname': 'none',
        #     'itemcount': 0,
        #     'is_authenticated': request.user.is_authenticated,
        #     'guest': guest,
        #     'description': 'Please fill the form below' if form.description == None else form.description


        # }
        # return render(request, 'regform.html', context)

""" when a user registers for a form, this function is used to save the values provided to the database
    also used to save the new values of a formvalue when user edits the form values""" 
def saveValue(request):
    if request.method == 'POST':
        formcode = request.POST.get('formcode')
        values = request.POST.get('values')
        user = request.user.username
        edit_mode = request.POST.get('editmode')

        if edit_mode== 'true':
            value_id = request.POST['valueid']
            if FormValue.objects.filter(id = value_id).exists():
                form_value = FormValue.objects.get(id =value_id)
                form_value.field_values = values
                form_value.save()
                context= {
                    'status' : 'success',
                    'message': 'Entries modified successfully'
                }
                return JsonResponse(context)
            else:
                context= {
                    'status' : 'failed',
                    'message': 'Value id does not exist'
                }
                return JsonResponse(context)
        else:
            if Form.objects.filter(form_code = formcode).exists():
                form = Form.objects.get(form_code=formcode)
                form_value = FormValue.objects.create(form_code=formcode, form_name= form.form_name, user_name= user, field_values=values)
                form_value.save()
                context= {
                    'status' : 'success',
                    'message': 'Registeration successfull'
                }
                return JsonResponse(context)
            else:
                context= {
                    'status' : 'failed',
                    'message': 'Form code does not exist'
                }
                return JsonResponse(context)

@login_required(login_url= 'login')
def publish(request):
    current_site = get_current_site(request)
    if request.method == 'POST':
        formname = request.POST.get('form-name')
        document = request.FILES.get('document')
        formcode = request.POST.get('form_code')
        formdata = request.POST.get('form_data')
        description = request.POST.get('description')
        description= None if description== '' else description
        formauthor = request.user.username

        if document is not None:#checking for the correct fie format
            _, fileextension =os.path.splitext(document.name)
            allowed_extensions = ['.doc', '.docx']
            if fileextension not in allowed_extensions:
                context = {
                    'status': 'failed',
                    'message': f'A {fileextension} file format is not allowed'
                }
                return JsonResponse(context)
            
        #checking if the form name already exits
        if Form.objects.filter(form_name=formname).exists():
            context = {
                'status': 'failed',
                'message': 'Form name already exists'
            }
            return JsonResponse(context)
        else:
            if Form.objects.filter(form_code = formcode).exists():
                context = {
                    'status': 'failed',
                    'message': 'Oops an error occurred, please try again'
                }
            else:
                if document is not None:
                    #checking if a dg table exists in the document
                    dgTable = checkForDgTable(document=document)
                    
                    if dgTable == True:
                        old_file_name = document.name
                        _, ext = os.path.splitext(old_file_name)                
                        document.name = 'doc-' + formcode + ext
                        new_form = Form.objects.create(form_name=formname, form_code=formcode, form_author=formauthor, appended_document=document, fields=formdata, original_doc_name = old_file_name, description = description)
                        new_form.save()
                        
                        
                        context = {
                            'status' : 'success',
                            'formname': formname,
                            'formcode': formcode,
                            'formauthor': formauthor,
                            'domain': current_site.domain
                        }
                        return JsonResponse(context)
                    else:
                        context = {
                            'status': 'failed',
                            'message': 'No DG table was not found in your document please provide a document containing a DG table'
                        }
                    
                else: 
                    new_form = Form.objects.create(form_name=formname, form_code=formcode, form_author=formauthor, fields=formdata, description=description)
                    new_form.save()
                    
                    context = {
                        'status' : 'success',
                        'formname': formname,
                        'formcode': formcode,
                        'formauthor': formauthor,
                        'domain': current_site.domain
                    }
            
                return JsonResponse(context)
    else:
        user = User.objects.get(id = request.user.id)
        username = user.username
        user_initials = ''.join(word[0] for word in username.split()[:2]).upper()
        context = {
            'domain': current_site,
            'user': user,
            'user_initials': user_initials
            
        }
        return render(request, 'publish.html', context)

#navigates to the viewform.html file 
@login_required(login_url='login')
def viewForm(request, formcode):
    current_site = get_current_site(request)
    if request.method == 'POST':#handle post request
        form = Form.objects.get(form_code= formcode)
        formValues = FormValue.objects.filter(form_code=formcode)

        context = {
            'form': form,
            'formvalues': list(formValues.values()),
            'domain': current_site
        }
        return JsonResponse(context)
    else:#handle get request
        try:
            form = Form.objects.get(form_code= formcode)
        except ObjectDoesNotExist:
            return render(request, 'form-not-found.html')
            
        formValues = FormValue.objects.filter(form_code=formcode)
        peopleCount = len(formValues)
        context = {
            'form': form,
            'formvalues': list(formValues.values()),
            'peoplecount': peopleCount,
            'domain': current_site
        }

        return render(request, 'viewform.html', context)
    
@login_required(login_url= 'login')
def getFormAndValues(request, formcode):
        if request.method == 'GET': 
            form = Form.objects.get(form_code=formcode)
            
            formValues = FormValue.objects.filter(form_code=formcode)
            serialized_form = serialize('json', [form])
           
            context = {
                'form': serialized_form,
                'formvalues': list(formValues.values())
            }
            
            return JsonResponse(context)
        
def modifyForm(request):
    if request.method == 'POST':
        new_name = request.POST.get('formname')
        new_document = request.FILES.get('document')
        formcode = request.POST.get('formcode')
        status = request.POST.get('status')
        description = request.POST.get('description')
        fileop = request.POST.get('fileop')


        if Form.objects.filter(form_code = formcode).exists():            
            form = Form.objects.get(form_code = formcode)

            
            if Form.objects.filter(form_name=new_name).exists() and form.form_name != new_name:
                return JsonResponse({
                    'statuscode': 400,
                    'message': 'A form with that name already exits'
                })

            if fileop == 'none':#User did not make changes to the document
                pass
            elif fileop == 'removed':#user removed the document
                #Delete the appended document
                doc_path = form.appended_document.path
                if os.path.exists(doc_path):
                    os.remove(doc_path)
                    print('deleted the appended document')
                
                form.appended_document = 'documents/sampledoc.docx'
                form.original_doc_name = 'document_name'
            else:#user selected a new document
                dg_table = checkForDgTable(new_document)
                if(dg_table):
                    original_file_name = new_document.name
                    new_document.name = form.appended_document.name

                    old_doucument_path = form.appended_document.path

                    #deleting the old document before saving the new one 
                    if form.appended_document.name == 'sampledoc.docx' or form.appended_document.name == 'documents/sampledoc.docx':
                         #----------GENERATING A NEW DOCUMENT---------------
                        new_doc_path = settings.DOC_DIR + 'doc-' + form.form_code + '.docx'
                        with open(new_doc_path, 'wb+') as destination:
                            for chunk in new_document.chunks():
                                destination.write(chunk)

                        # with open(new_doc_path, 'rb') as new_file:
                        #     form.appended_document.save(f'doc-{form.form_code}.docx', File(new_file), save=True)

                        form.appended_document = new_doc_path
                    else:
                        if os.path.exists(old_doucument_path):
                            os.remove(old_doucument_path)
                        
                        form.appended_document = new_document

                    form.original_doc_name = original_file_name
                        
                else:
                    return JsonResponse({
                        'statuscode': 400,
                        'message': 'Updated document does not contain a dg table'
                    }, status = 400)

            if form.form_name != new_name:#checking if form name needs to be updated
                form.form_name = new_name
                form_values= FormValue.objects.filter(form_code = form.form_code)#changing all the names for the form values
                #TODO: make form values a forien key
                for form_value in form_values:
                    form_value.form_name = new_name
                    form_value.save()
            
            #update the status and the description
            form.status = status
            form.description = description

            form.save()
            print(form)
            return JsonResponse({
                'statuscode': 200,
                'message': 'Form updated succesfully'
            }, status=200)
        else:
            return JsonResponse({
                'statuscode': 404,
                'message': 'Form not found'
                }, status=400)
    else:
        return JsonResponse({
            'statuscode': 400, 
            'message': 'GET method not allowed'}, status=400)        

def editForm(request):
    if request.method == 'POST':
        new_name = request.POST.get('formname')
        new_document = request.FILES.get('document')
        formcode = request.POST.get('formcode')

        #checking if the form with the provided form code exists
        if Form.objects.filter(form_code= formcode).exists():
            form = Form.objects.get(form_code= formcode)
       
            if new_document is not None:
                #this means the user selected a document
                #chechking for a dg table in the new document provided 
                dgTable = checkForDgTable(new_document)
                if dgTable == True:
                    original_file_name = new_document.name
                    new_document.name = form.appended_document.name
                    
                    old_doucument_path = form.appended_document.path

                    #deleting the old document before saving the new one 
                    if form.appended_document.name == 'sampledoc.docx':
                        #----------GENERATING A NEW DOCUMENT---------------
                        new_doc = settings.DOC_DIR + 'doc-' + form.form_code + '.docx'
                        with open(new_doc, 'wb+') as destination:
                            shutil.copyfileobj(new_document, destination)
                        
                        document = Document(new_doc).save()
                        form.appended_document = document

                    else:
                        if os.path.exists(old_doucument_path):
                            os.remove(old_doucument_path)
                            form.appended_document = new_document

                    form.original_doc_name = original_file_name


                    form.form_name = new_name #changing the name of the form
                    #changing the form names for every formvalue object
                    form_values= FormValue.objects.filter(form_code = form.form_code)
                    for form_value in form_values:
                        form_value.form_name = new_name
                        form_value.save()
                else:
                    return JsonResponse({'status': 'failed', 'message': 'No dg table in appended document'})
            else:
                #this means user did not select any doument
                form.form_name = new_name

                #changing the form names for every formvalue object
                form_values= FormValue.objects.filter(form_code = form.form_code)
                for form_value in form_values:
                    form_value.form_name = new_name
                    form_value.save()
              
            form.save()
            context = {
                'status': 'success',
                'message': 'Form successfully modified'
            }
            return JsonResponse(context)
        else:
            context= {
                'status': 'failed',
                'message': str(formcode) + ' not found'
            }

            return JsonResponse(context)
    else:
        pass

#Returns true or false depending on if the provided document contains a dg table or not 
def checkForDgTable(document):
    doc = Document(document)
    dgTable = False
    tables = doc.tables
    for table in tables:
        rowcount = len(table.rows)#getting how many rows are in the table
        if rowcount> 1:
            pass
        else:
            row = table.rows[0].cells
            if len(row) >1:
                pass
            else:
                cellValue = row[0].text
                if cellValue.lower() == 'dg':
                    dgTable = True
                else:
                    pass
    return dgTable

#deletes temporary files 
def deleteTemp(request, formcode):
    if Form.objects.filter(form_code= formcode).exists():
        form = Form.objects.get(form_code =formcode)
        

        file_name = 'temp-doc-' + form.form_code + '.docx'
        if os.path.exists(file_name):
            os.remove(file_name)
            context = {
                'status' : 'success',
                'message': 'file deleted successfully'
            }
            return JsonResponse(context)
        else:
            context = {
                'status' : 'success',
                'message': 'file deleted successfully'
            }
            return JsonResponse(context)
    else:
        context = {
            'status': 'failed',
             'message': 'Form not found'
        }

        return JsonResponse(context)
    
#delets form with provided form code 
def deleteForm(request, formcode):

    if Form.objects.filter(form_code = formcode).exists():
        form = Form.objects.get(form_code = formcode)

        #deleting all the formvalues associated with this formcode 
        formvalues = FormValue.objects.filter(form_code= formcode)
        if len(formvalues) > 0:
            for value in formvalues:
                value.delete()

        temp_doc = settings.TEMP_DIR + 'temp-doc-' + form.form_code + '.docx'
        if os.path.exists(temp_doc):
            os.remove(temp_doc)

        document_path = settings.MEDIA_ROOT +'/documents/doc-' + form.form_code + '.docx'
        if os.path.exists(document_path):
            os.remove(document_path)

        form.delete()
        context = {
            'status': 'success',
             'message': 'Form deleted successfully'
        }
        return JsonResponse(context)
    else:
        context = {
            'status': 'failed',
             'message': 'Form not found'
        }
        return JsonResponse(context)
    
#generates a temporary document and downloads it 
def generateDoc(request, formcode, docname, numbered, alph, factor, transverse):
   temp_dir = settings.TEMP_DIR
   if request.method == 'GET': 
        #checking if form exists 
        if Form.objects.filter(form_code=formcode).exists:
            form = Form.objects.get(form_code= formcode)
            formValues = FormValue.objects.filter(form_code=formcode)
            
            is_numbered = {'true': True, 'false': False}.get(numbered.lower())   
            is_alphabetical_ordered = {'true': True, 'false': False}.get(alph.lower())
            
            #----------GENERATING A TEMPORARY DOCUMENT---------------
            temp_doc = temp_dir + 'temp-doc-' + form.form_code + '.docx'
            print('appended', form.appended_document)

            with open(temp_doc, 'wb+') as destination:
                with form.appended_document.open('rb') as source: 
                    shutil.copyfileobj(source, destination)
                
            #getting the number of rows and number of columns required to make a table out of the form values
            #adding 1 to the length of my form values for the table to account for the headers
            num_rows = len(list(formValues.values())) +1
            if is_numbered:
                #adding 1 column to account for the number colomn if list is supposed to be numbered
                num_cols =  len(json.loads(form.fields)) +1
            else:
                num_cols =  len(json.loads(form.fields))
           
            doc = Document(temp_doc)
            tables = doc.tables
            for table in tables:
                rowcount = len(table.rows)
                if rowcount> 1:
                    pass
                else:
                    row = table.rows[0].cells
                    if len(row) >1:
                        pass
                    else:
                        cellValue = row[0].text
                        if cellValue.lower() == 'dg':
                            #we found a dg table!
                            parent_table = table._element.getparent()
                            new_table = doc.add_table(rows=num_rows, cols=num_cols)
                            new_table.style = 'Table Grid'
                            #inserting our table right below the dg table 
                            parent_table.insert(parent_table.index(table._element) +1, new_table._element)
                            #removing the dg table leavinig only the new one 
                            parent_table.remove(table._element)
                        else:
                            pass
                
            #setting the content of the header cells
            header_cells = new_table.rows[0].cells
            if is_numbered:
                header_cells[0].text = ''
                hIndex = 1
            else:
                 hIndex =0
            for header in json.loads(form.fields):
                cell = header_cells[hIndex]
                if header.get('name') == '-':
                
                    cell.text = ''
                else:
                    cell.text = header.get('name')

                if cell.text != '':
                    cell.paragraphs[0].runs[0].bold = True
                hIndex +=1

            form_value_list = list(formValues.values())
            
            if is_alphabetical_ordered:
                if transverse == 'asc':
                    form_value_list.sort(key= partial(sort_by_field_values, factor=factor), reverse=False)
                else:
                    form_value_list.sort(key= partial(sort_by_field_values, factor=factor), reverse=True)


            #setting the other rows of the table 
            #iterate starting from row index 1 to row index num_rows, this is done to exclude the headers
            for i in range(1, num_rows):
                row_cells = new_table.rows[i].cells
                form_fields = json.loads(form.fields)
                val = form_value_list[i-1]
                pIndex = 0
                for index, cell in enumerate(row_cells):
                    if is_numbered:
                        if index == 0:
                            cell.text = str(i)
                        else:
                            field = form_fields[pIndex]
                            if field.get('datatype') == 'empty':
                                cell.text = ''
                            else:
                                cell.text = field.get('name')
                                for item in json.loads(val.get('field_values')):
                                    if item.get('name') == field.get('name'):
                                        cell.text = item.get('value')
                            pIndex +=1
                    else:
                        field = form_fields[pIndex]
                        if field.get('datatype') == 'empty':
                            cell.text = ''
                        else:
                            cell.text = field.get('name')
                            for item in json.loads(val.get('field_values')):
                                if item.get('name') == field.get('name'):
                                    cell.text = item.get('value')
                        
                        pIndex +=1

            
            doc.save(temp_doc)

            #-----DOWNLOADING THE GENERATED TEMPORARY DOCUMENT----------
            file_name = temp_dir+ 'temp-doc-' + form.form_code + '.docx'
       
            if os.path.exists(file_name):
                file = open(file_name, 'rb')
                
                response = FileResponse(file)
                response['Content-disposition'] ='attachment; filename= "' + docname + '"'
                return response
            else:
                context = {
                    'status': 'failed',
                    'message': 'Document does not exist'
                }
                return JsonResponse(context, status=500)
        else:
            context = {
                'status': 'failed',
                'message': 'Form not found'
            }
            return JsonResponse(context, status= 500)
        
def sort_by_field_values(item, factor):
    for item in json.loads(item.get('field_values')):
        if item.get('name') == factor:
            return item.get('value')
        
def sendEmail(request):
    if request.method == 'POST':
        
        form = mForms.ContactForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data['email']
            name = form.cleaned_data['name']
            subject = form.cleaned_data['subject']
            message = form.cleaned_data['message']
           
            contact = Contact.objects.create(email=email, name= name, subject= subject, message = message)
            contact.save()

            html = render_to_string('emails/contactform.html', {
                'name': name,
                 'email': email,
                 'subject':subject,
                 'message': message
            })

            sender = 'Data gridder <' + str(settings.EMAIL_HOST_USER) + '>' 
            send_mail(subject, message, sender, ['support@datagridder.com', 'onuhudoudo@gmail.com'],html_message=html,fail_silently=False)
            context = {
                'status': 'success',
                'statusCode': 200,
            }
            return JsonResponse(context, status=200)
        else:
            context = {
                'status': 'failed',
                'statusCode': 400,
            }
            return JsonResponse(context, status=400)

        #redirecting the user to the previous page they were in 
        # return redirect(request.META['HTTP_REFERER']+ '?modal=sent_mail')


"""this deletes a form value with the id provided in the url"""
def deleteEntry(request, entry_id):
    entry_id = int(entry_id)
    if FormValue.objects.filter(id= entry_id).exists():
        entry = FormValue.objects.get(id= entry_id)
        entry.delete()
        context = {
            'status': 'success',
            'message': 'Entry deleted successfully',
            'statuscode': 200
        }
        return JsonResponse(context, status=200)
    else:
        context= {
            'status': 'failed',
            'message': 'Entry does not exist',
            'statuscode': 404
        }
        return JsonResponse(context, status=404)

def documentation(request):
    if request.user.is_authenticated:
        user_object = User.objects.get(id= request.user.id)
        user_profile = Profile.objects.get(user=user_object)
        context= {
            'user_exists': 'true',
            'user_profile' : user_profile
        }
        return render(request, 'documentation.html', context)
    else:
        context = {
            'user_exists': 'false'
        }
        return render(request, 'documentation.html', context)
        
def showCurrentSite(request):
    current_site = get_current_site(request)
    return redirect('/')

def contactUs(request):
    contact_form = mForms.ContactForm()
    context = {
        'contact_form': contact_form
    }
    return render(request,'contactus.html', context)