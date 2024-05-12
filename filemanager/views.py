from django.shortcuts import render
from docx import Document
from django.http import JsonResponse
from django.views.generic import View
import os


# Create your views here.



class validateExistingDocument(View):
    def post(self, request):
        request_document = request.FILES.get('document')
        _, file_extension = os.path.splitext(request_document.name)
        allowed_extensions = ['.docx', '.doc']
        context = {}
        success_message = 'Hurray! a dg column was found in your document, you can now proceed to publishing your form'
        if file_extension.lower() in allowed_extensions:
            document = Document(request_document)
            tables = document.tables
            for table in tables:
                rowcount= len(table.rows)
                if rowcount >1:#this means there are multiple rows in the table
                    for row in table.rows:
                        columns = row.cells
                        for cell in columns:
                            cellvalue = cell.text
                            if cellvalue.lower() == 'dg':#we found a dg table
                                context = {
                                    'status': 'success',
                                    'message': success_message,
                                    'statusCode': 200
                                }
                                print('found a dg table')
                                return JsonResponse(context)            
                else:
                    row = table.rows[0].cells
                    if len(row) >1:#this means there are multiple columns in a row
                        for column in row:
                            cellvalue = column.text
                            if cellvalue.lower() == 'dg':#we found a dg column
                                context = {
                                'status': 'success',
                                'message': success_message,
                                'statusCode': 200
                            }
                            print('found a dg table')
                            return JsonResponse(context)
                    else:
                        cellvalue = row[0].text
                        if cellvalue.lower() =='dg': #we found a dg column
                            context = {
                                'status': 'success',
                                'message': success_message,
                                'statusCode': 200
                            }
                            print('found a dg table')
                            return JsonResponse(context)
                        else:
                            pass

            context = {
                'status': 'failed',
                'statusCode': 401,
                'message': 'There was no dg column found in your document, please ensure to include a new row with at leeast one column having the text dg within it '
            }
            return JsonResponse(context, status=401)
        else:
            context ={
                'status': 'failed',
                'statusCode': 400,
                'message' : f'A {file_extension} document format is not allowed',
            }

            return JsonResponse(context, status=400)
